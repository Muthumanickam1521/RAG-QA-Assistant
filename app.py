from indexer import doc_to_index, delete_index
from retriever import index_to_response

import time
import os
from PyPDF2 import PdfReader
import numpy as np
import pandas as pd

from llama_index.core import Document
from llama_index.llms.gemini import Gemini
from llama_index.readers.web import BeautifulSoupWebReader

import streamlit as st
from streamlit import session_state as ss

def read_pdf(uploaded_file):
    print("Reading PDF file...")
    pdf_reader = PdfReader(uploaded_file)
    texts = convert_pdf_to_text(pdf_reader)
    return convert_text_to_doc(texts)

def read_webpage(url):
    print("Reading webpage...")
    web_reader = BeautifulSoupWebReader
    return web_reader.load_data(urls=[url])

def read_csv(uploaded_file):
    print("Reading CSV file...")
    try:
        df = pd.read_csv(uploaded_file)
    except UnicodeDecodeError:
        print("UTF-8 decoding failed, trying ISO-8859-1 encoding.")
        df = pd.read_csv(uploaded_file, encoding="ISO-8859-1")
    text = df.to_string()
    return convert_text_to_doc(text)

def read_excel(uploaded_file):
    print("Reading Excel file...")
    df = pd.read_excel(uploaded_file)
    text = df.to_string()
    return convert_text_to_doc(text)

def read_docx(uploaded_file):
    print("Reading DOCX file...")
    from docx import Document
    doc = Document(uploaded_file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return convert_text_to_doc(text)

def read_text(text_file):
    text = text_file.read().decode("utf-8")
    return convert_text_to_doc(text)

def convert_text_to_doc(texts):
    return Document(text=texts)

def convert_pdf_to_text(pdf_reader):
    texts = ""
    for page in pdf_reader.pages:
        texts += page.extract_text()
    return texts

GOOGLE_API_KEY = "AIzaSyCcz5K_IEIq_cW_2Y3hagkkDqr_3cPIpx8"
PINECONE_API_KEY = "pcsk_3HFKTd_R36Vrr5AoFVURe4AP1Ez76UMq11Cnwwm8t6Zhb19ZqSa9FYR8fwAiBxAXdyHWKP"
os.environ["PINECONE_API_KEY"] = PINECONE_API_KEY
os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

response_lang = "English"
bot_name = "Smart Bot"
bot_tone = "friendly, funny, informal, & helpful"
bot_avator_filename = "avatar.jpg"
bot_avator_filepath = os.path.join(os.getcwd(), "asset", bot_avator_filename)

model_name = "models/embedding-001"
top_p = 0.95
max_tokens = 200
default_temperature = 0.4

documents = None
uploaded_file = None

st.title("Q/A RAG Bot with Gemini 2.0 Flash")

if 'file' not in ss:
    ss.file = None
    ss.messages = []
    ss.indexed = False 

if len(ss.messages) == 0:
    delete_index()
    bot_avator_filepath = os.path.join(os.getcwd(), "asset", "avatar.jpg")
    ss.messages.append({
        "role": "ai", 
        "content": "Hi!\n\n I am Smart Bot. I'm excited to answer your questions about a file. Feel free to upload one.", 
        "avatar": bot_avator_filepath
    })

with st.sidebar:
    uploaded_file = st.file_uploader("Upload a file (PDF, DOCX, CSV, Excel, Text, or URL)", type=["pdf", "docx", "csv", "xlsx", "txt"])

documents = None
if uploaded_file is not None and not ss.indexed:
    file_extension = uploaded_file.name.split('.')[-1].lower()

    if file_extension == "txt":
        documents = read_text(uploaded_file)
    elif file_extension == "docx":
        documents = read_docx(uploaded_file)
    elif file_extension == "pdf":
        documents = read_pdf(uploaded_file)
    elif file_extension == "csv":
        documents = read_csv(uploaded_file)
    elif file_extension == "xlsx":
        documents = read_excel(uploaded_file)
    else:
        url = st.text_input("Provide URL to fetch data from")
        if url:
            documents = read_webpage(url)

    if documents:
        flag_indexed = doc_to_index(documents)
        if flag_indexed:
            ss.indexed = True
            ss.documents = documents
        else:
            ss.indexed = False

with st.sidebar:
    temperature = st.sidebar.select_slider("Temperature", options=[0, 0.2, 0.4, 0.6, 0.8, 1], value=default_temperature)
    ss.temperature = temperature
    default_temperature = temperature

    response_lang = st.sidebar.selectbox("Response Language", options=["English", "Tamil", "Hindi", "French", "German"])
    ss.response_lang = response_lang


prompt = f"""
You are {bot_name} and you need to assist people by answering their questions.
You are an intellectual and helpful Question Answering assistant chatbot. 
Your objective is to answer the query about references. 

General instructions to Responding:
1) Never respond with words like based on, according to, snippet, reference, section, or any words related to provided context. 
2) Use bullet points or number list when necessary. 
3) Use latex typed equations or even texts.
4) Always provide a response in {response_lang} language with tone {bot_tone}.     
5) Always provide a response in markdown format with proper markdown syntax. Overall the response should look neat and easy to understand.
6) Your target is general public. So if needed explain things in simple terms.
7) Respond in less words unless it's absolute necessary or it would be conveying good.

Role:
You are working as a Senior Full Stack Data Scientist at a product company. 
You are a detective in finding insights and putting them in ley-man terms.
You are also a Physicist and you can explain complex concepts in easy terms.

"""

for message in ss.messages:
    if message["role"] == "human":
        st.chat_message("human").markdown(message["content"])
    else:
        st.chat_message("ai", avatar=bot_avator_filepath).markdown(message["content"])

if "indexed" in ss:
    if query := st.chat_input("Ask anything about the file"):
        with st.chat_message("human"):
            st.markdown(query)
        ss.messages.append({"role": "human", "content": query})

        references = index_to_response(query)

        for idx, reference in enumerate(references, 1):
            prompt += f"""Reference {idx}: {reference.get_content()}\n\n"""
        prompt += f"""Question: {query}\n\n"""
        
        llm = Gemini(model="models/gemini-2.0-flash", temperature=temperature, max_tokens=max_tokens, top_p=top_p)
        response = llm.complete(prompt)

        with st.chat_message("ai", avatar=bot_avator_filepath):
            placeholder = st.empty() 
            response_text = " "
            for word in response.text.split(" "):
                response_text += word + " "
                placeholder.markdown(response_text)  
                time.sleep(0.15) 
        ss.messages.append({"role": "ai", "content": response_text, "avatar": bot_avator_filepath})
